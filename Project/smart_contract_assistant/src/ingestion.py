
import pathlib
import pdfplumber
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.runnables import RunnableLambda
from ..config import CHUNK_SIZE, CHUNK_OVERLAP

def extract_text(file_path: str) -> str:
    ext = pathlib.Path(file_path).suffix.lower()
    if ext == ".pdf":
        pages = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
        return "\n\n".join(pages)
    elif ext == ".docx":
        doc = DocxDocument(file_path)
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        raise ValueError(f"Unsupported file: {ext}")


def ingest_document(file_path: str) -> list:
    """
    LCEL Ingestion Pipeline:
    raw_text → splitter | Document wrapper
    """
    filename = pathlib.Path(file_path).name
    raw_text = extract_text(file_path)
    if not raw_text.strip():
        raise ValueError(f"No text extracted from '{filename}'.")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", " ", ""],
    )

    split_fn   = RunnableLambda(lambda text: splitter.split_text(text))
    wrap_fn    = RunnableLambda(lambda chunks: [
        Document(page_content=c, metadata={"source": filename, "chunk_index": i})
        for i, c in enumerate(chunks)
    ])

    ingestion_chain = split_fn | wrap_fn
    docs = ingestion_chain.invoke(raw_text)

    print(f"[Ingestion] '{filename}' → {len(docs)} chunks")
    return docs
